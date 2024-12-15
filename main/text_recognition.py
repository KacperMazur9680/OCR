from transformers import TrOCRProcessor, VisionEncoderDecoderModel
from PIL import Image
import os
import re
from docx import Document
from reportlab.pdfgen import canvas
import json

class TrOCRProcessorTool:
    def __init__(self, model_name="microsoft/trocr-large-printed", roi_dir="saved_rois", output_word_file="ocr_results.docx"):
        self.processor = TrOCRProcessor.from_pretrained(model_name, use_auth_token=False)
        self.model = VisionEncoderDecoderModel.from_pretrained(model_name, use_auth_token=False)
        self.roi_dir = roi_dir
        self.output_word_file = output_word_file

    def extract_number(self, filename):
        """Extract numeric parts from filenames for sorting."""
        match = re.search(r'\d+', filename)
        return int(match.group()) if match else float('inf')

    def process_rois(self, save_word=False):
        """
        Process ROI images in a single directory, perform OCR, and optionally save results to a Word file.
        
        Args:
            save_word (bool): Whether to save the OCR results to a Word document.
        
        Returns:
            dict: OCR results mapped to image names.
        """
        if not os.path.exists(self.roi_dir):
            print(f"Error: ROI directory '{self.roi_dir}' does not exist.")
            return {}

        ocr_results = {}

        # Filter only image files in the directory
        valid_extensions = {".png", ".jpg", ".jpeg", ".bmp", ".tiff"}
        sorted_files = sorted(
            [f for f in os.listdir(self.roi_dir) if os.path.splitext(f)[1].lower() in valid_extensions],
            key=self.extract_number
        )

        for img_name in sorted_files:
            img_path = os.path.join(self.roi_dir, img_name)

            if os.path.isfile(img_path):  # Ensure it's a valid file
                text_output = self.perform_ocr(img_path)
                ocr_results[img_name] = text_output

        # Save to Word only if explicitly requested
        if save_word:
            self.save_to_word(ocr_results)

        return ocr_results

    def perform_ocr(self, img_path):
        """
        Perform OCR on a single image.
        Args:
            img_path (str): Path to the image file.
        Returns:
            str: Recognized text.
        """
        image = Image.open(img_path).convert("RGB")
        pixel_values = self.processor(images=image, return_tensors="pt").pixel_values

        # Generate the text output
        generated_ids = self.model.generate(pixel_values, max_new_tokens=50)  # Adjust max_new_tokens if needed
        text_output = self.processor.batch_decode(generated_ids, skip_special_tokens=True)[0].lower()

        return text_output

    def save_to_word(self, ocr_results):
        """
        Save OCR results to a Word document.
        Args:
            ocr_results (dict): Dictionary containing recognized text.
        """
        doc = Document()
        doc.add_heading("OCR Results", level=1)

        for roi_name, text in ocr_results.items():
            doc.add_heading(f"ROI: {roi_name}", level=2)
            doc.add_paragraph(text)

        doc.save(self.output_word_file)
        print(f"OCR results saved to {self.output_word_file}")

    def save_to_pdf_with_layout(self, image_path, bboxes, ocr_results, output_pdf="ocr_results.pdf"):
        """
        Save OCR results to a PDF with dynamically adjusted font size based on bounding box height.
        Args:
            image_path (str): Path to the original image.
            bboxes (list): List of bounding boxes [(x_min, y_min, x_max, y_max), ...].
            ocr_results (dict): Recognized text for each ROI, keyed by ROI filenames.
            output_pdf (str): Path to the output PDF file.
        """
        # Load saved coordinates to ensure correct order
        coordinates_file = os.path.join("saved_coordinates.json")
        if os.path.exists(coordinates_file):
            with open(coordinates_file, "r") as f:
                bboxes = json.load(f)

        # Open the image to get dimensions
        image = Image.open(image_path)
        image_width, image_height = image.size

        # Create a PDF canvas
        pdf = canvas.Canvas(output_pdf, pagesize=(image_width, image_height))

        # Draw the original image as the background
        pdf.drawImage(image_path, 0, 0, width=image_width, height=image_height)

        # Combine bounding boxes and OCR results
        rois_with_text = list(zip(bboxes, list(ocr_results.values())))

        # Sort bounding boxes top-to-bottom, then left-to-right
        def sort_key(roi):
            x_min, y_min, x_max, y_max = roi[0]
            return (y_min, x_min)

        rois_with_text.sort(key=sort_key)

        # Add recognized text to the PDF based on the sorted bounding boxes
        for (x_min, y_min, x_max, y_max), text in rois_with_text:
            y_min_pdf = image_height - y_min  # Convert y-coordinates

            bbox_height = y_max - y_min
            font_size = max(6, bbox_height * 0.6)  # Minimum font size of 6

            pdf.setFont("Helvetica-Bold", font_size)  # Use bold font
            pdf.setFillColorRGB(1, 0, 0)  # Set font color to red
            pdf.drawString(x_min, y_min_pdf - font_size, text)

        # Save the PDF
        pdf.save()
        print(f"OCR results with layout and dynamic font size saved to {output_pdf}")



